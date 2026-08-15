#!/usr/bin/env python3
"""Convert SOP markdown files to .docx Word documents."""

import re
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def parse_inline(paragraph, text):
    """Parse inline markdown (bold, code, etc.) and add runs to paragraph."""
    # Pattern: **bold**, `code`, *italic*
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC8, 0xA9, 0x7E)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_table(doc, rows_data):
    """Add a formatted table to the document."""
    if not rows_data:
        return

    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if i == 0:
                # Header row
                set_cell_shading(cell, '2D2D2D')
                run = p.add_run(cell_text.strip())
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                parse_inline(p, cell_text.strip())
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def md_to_docx(md_path, docx_path):
    """Convert a single markdown file to docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                p.paragraph_format.left_indent = Cm(1)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Headers
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0xC8, 0xA9, 0x7E)
            i += 1
            continue
        if stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0xC8, 0xA9, 0x7E)
            i += 1
            continue
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
            i += 1
            continue

        # Blockquote (> text)
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            parse_inline(p, text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            i += 1
            continue

        # Table
        if '|' in stripped and stripped.startswith('|'):
            table_rows = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            add_table(doc, table_rows)
            continue

        # Checklist
        if stripped.startswith('- [ ] ') or stripped.startswith('- [x] '):
            checked = '[x]' in stripped
            text = stripped[6:]
            p = doc.add_paragraph()
            symbol = '☑' if checked else '☐'
            run = p.add_run(f'{symbol}  ')
            run.font.size = Pt(10)
            parse_inline(p, text)
            p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(f'{m.group(1)}.  ')
            run.bold = True
            parse_inline(p, m.group(2))
            p.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    return docx_path


def main():
    sop_dir = Path('/Users/mac/mekong-cli/FnB-Container-Caffe/docs/sop')
    output_dir = sop_dir  # Same directory

    md_files = sorted(sop_dir.glob('*.md'))
    md_files = [f for f in md_files if f.name != 'README.md']

    print(f'Converting {len(md_files)} SOP files to .docx...\n')

    for md_file in md_files:
        docx_file = output_dir / md_file.with_suffix('.docx').name
        try:
            result = md_to_docx(str(md_file), str(docx_file))
            print(f'  ✅ {md_file.name} → {docx_file.name}')
        except Exception as e:
            print(f'  ❌ {md_file.name} → ERROR: {e}')

    print(f'\nDone! Files saved to: {output_dir}')


if __name__ == '__main__':
    main()
